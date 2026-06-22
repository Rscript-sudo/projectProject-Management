#!/usr/bin/env python3
"""
format_docx.py - DOCX文档统一格式刷工具（v4.3.0）
读取 ~/.claude/skills/共享格式规范/DOCX格式规范.json，
自动为所有 docx 应用全局统一格式规范。

v4.3.0 更新：
- 移除 should_insert_page_break 中未实现的附录分页逻辑（死代码）
- 顶部统一导入 etree，移除函数内重复导入
- 字体检测跳过非 Linux 系统（macOS 不支持 fc-list）

v4.2.0 更新：
- 改进分页逻辑：支持显式分页标记（---分页---）
- 改进编号项判定：正则匹配优先，缩进作为辅助
- 新增 body_no_indent 类型：用于统计说明等无缩进正文
- 改进空段落处理：保持完全空白

v4.1.0 更新：
- 完善单位说明（所有数值统一用pt，EMU转换在内部处理）
- 添加 page_break 分页控制支持
- 统一字体名称（仿宋_GB2312）
- 改进字体 fallback 机制
- 修复分页符设置逻辑

v4.0.0 更新：
- 移除 profile 机制，使用全局统一规范
- 规范文件集中存放在 ~/.claude/skills/共享格式规范/
- 支持 --dry-run 检测、--light 轻量模式（仅修正字体/字号/行距，不重分类）
- 字体 fallback 机制：系统无指定字体时自动降级

用法：
    python format_docx.py <input.docx> [output.docx] [--dry-run] [--light]
    python format_docx.py <input.docx>  # 就地修改
    python format_docx.py <input.docx> --dry-run  # 仅检测，不写入
    python format_docx.py <input.docx> --light   # 轻量模式（有模板文档推荐）
"""
import sys
import os
import re
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# ──────────────────────────────────────────────────────────────
# 路径解析
# ──────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent.resolve()
# 格式规范 JSON 在同一目录下
FORMAT_SPEC_PATH = SCRIPT_DIR / "DOCX格式规范.json"

# EMU 转换常量（内部使用，外部接口统一用pt）
PT_TO_EMU = 12700
CM_TO_EMU = 360000

def pt(num):
    """将pt值转换为EMU"""
    return int(num * PT_TO_EMU)

def cm(num):
    """将cm值转换为EMU"""
    return int(num * CM_TO_EMU)


# ──────────────────────────────────────────────────────────────
# 规范加载
# ──────────────────────────────────────────────────────────────
def load_format_spec():
    """加载全局统一格式规范。"""
    if not FORMAT_SPEC_PATH.exists():
        raise FileNotFoundError(
            f"格式规范文件不存在: {FORMAT_SPEC_PATH}\n"
            f"请确认格式规范文件存在于 templates/format-spec/ 目录下"
        )
    with open(FORMAT_SPEC_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


# ──────────────────────────────────────────────────────────────
# 分页符设置（统一使用 pageBreakBefore 属性）
# ──────────────────────────────────────────────────────────────
def set_page_break_before(para):
    """
    给段落设置分页符（pageBreakBefore）。
    这是Word文档分页的标准方式。
    """
    pPr = para._p.get_or_add_pPr()
    pageBreak = OxmlElement('w:pageBreakBefore')
    pageBreak.set(qn('w:val'), '1')
    pPr.append(pageBreak)


def is_page_break_marker(text, spec):
    """
    判断段落是否为显式分页标记。
    常见标记：---分页---、分页、pagebreak等
    """
    pb_config = spec.get("page_break", {})
    markers = pb_config.get("page_break_markers", ["---分页---", "分页", "pagebreak"])

    text_lower = text.strip().lower()
    for marker in markers:
        if marker.lower() in text_lower:
            return True
    return False


def should_insert_page_break(text, spec):
    """
    根据规范判断当前段落是否需要分页。
    显式分页标记：段落内容包含 page_break_markers
    """
    return is_page_break_marker(text, spec)


# ──────────────────────────────────────────────────────────────
# 类型检测
# ──────────────────────────────────────────────────────────────
PARA_IDX_COUNTER = [0]

def detect_type(para, spec):
    """
    根据段落文本内容判断段落类型。
    优先级从高到低：
    空段落 → 显式分页标记 → 落款 → 附录 → 封面 → 图表标题 → 注释 →
    标题层级（一二三级） → 编号项 → 无缩进正文 → 正文
    """
    global PARA_IDX_COUNTER
    para_idx = PARA_IDX_COUNTER[0]
    PARA_IDX_COUNTER[0] += 1

    text = para.text.strip()
    td = spec.get("type_detection", {})

    # 空段落（无文本）
    if not text:
        return "blank"

    # 显式分页标记
    if is_page_break_marker(text, spec):
        return "page_break_marker"

    # 落款关键字（优先级较高）
    for kw in td.get("closing_keywords", []):
        if kw in text:
            return "closing"

    # 附录关键字
    for kw in td.get("appendix_keywords", []):
        if kw in text:
            return "appendix_title"

    # 封面检测 - 单位信息
    for kw in td.get("cover_unit_keywords", []):
        if kw in text:
            return "cover_unit_info"

    # 封面检测 - 报告日期
    for kw in td.get("cover_report_date_keywords", []):
        if kw in text:
            return "cover_report_date"

    # 封面-标题检测（前 10 段）
    if para_idx < 10:
        ct_kws = td.get("cover_title_keywords", [])
        if ct_kws and len(text) >= 4 and len(text) < 30:
            if "（" not in text and "、" not in text:
                matched = sum(1 for kw in ct_kws if kw in text)
                if matched >= 2:
                    return "cover_title"

    # 图表标题检测
    chart_prefix = td.get("chart_title_prefix", "")
    if chart_prefix and re.match(chart_prefix, text):
        return "chart_title"

    # 注释检测
    note_prefix = td.get("note_prefix", "")
    if note_prefix and re.match(note_prefix, text):
        return "note"

    # 一级标题（一、二、三、……）
    h1_pattern = td.get("h1_pattern", "^\\s*[一二三四五六七八九十]+[、，。]")
    if re.match(h1_pattern, text):
        return "h1"

    # 二级标题（（一）（二）（三））
    h2_pattern = td.get("h2_pattern", "")
    if h2_pattern and re.match(h2_pattern, text):
        return "h2"

    # 三级标题（1. 2. 3. 或 1.1 2.2）
    h3_pattern = td.get("h3_pattern", "")
    if h3_pattern and re.match(h3_pattern, text):
        return "h3"

    # 三级标题检测（关键词匹配，备用）
    h3_kws = td.get("h3_keywords", [])
    if h3_kws and len(text) >= 8 and len(text) <= 60:
        ex_kws = td.get("h3_exclude_keywords", [])
        if any(kw in text for kw in h3_kws) and not any(kw in text for kw in ex_kws):
            return "h3"

    # 四级标题（（1）、（2））
    h4_pattern = td.get("h4_pattern", "")
    if h4_pattern and re.match(h4_pattern, text):
        return "h4"

    # 五级标题（①②③）
    h5_pattern = td.get("h5_pattern", "")
    if h5_pattern and re.match(h5_pattern, text):
        return "h5"

    # 编号项判定（正则匹配优先，缩进作为辅助验证）
    # 注意：之前的版本把缩进作为主要判定条件，导致很多正文被误判为编号项
    numbered_pattern = td.get("numbered_pattern", "^\\s*\\d+[.、]")
    if re.match(numbered_pattern, text):
        # 正则匹配成功，直接判定为编号项（不再依赖缩进）
        return "numbered"

    # 无缩进正文检测（统计说明等连续段落）
    # 条件：包含关键词且无首行缩进
    body_no_indent_kws = td.get("body_no_indent_keywords", [])
    if body_no_indent_kws and any(kw in text for kw in body_no_indent_kws):
        pf = para.paragraph_format
        # 无首行缩进时判定为 body_no_indent
        if not pf.first_line_indent or pf.first_line_indent == 0:
            return "body_no_indent"

    # 正文（默认）- 有首行缩进的正文
    return "body"


# ──────────────────────────────────────────────────────────────
# 字体设置（直接操作 XML，确保东亚字体正确应用）
# ──────────────────────────────────────────────────────────────
def set_run_font(run, font_name):
    """通过直接操作 XML 设置 run 的字体（含东亚字体）。"""
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(r, qn("w:rPr"))
    fonts = rPr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = etree.SubElement(rPr, qn("w:rFonts"))
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    run.font.name = font_name


def get_fallback_font(font_name, spec):
    """
    获取字体的 fallback 链中的第一个可用字体。
    """
    fallback = spec.get("font_fallback", {})
    fallback_chain = fallback.get(font_name, ["宋体"])
    return fallback_chain[0] if fallback_chain else "宋体"


def check_and_fix_font(run, allowed_fonts, spec, fallback_font="宋体"):
    """
    检查 run 的字体是否在 allowed_fonts 清单内，不在则替换。
    用于 light 模式。
    """
    current_font = run.font.name or ""
    if current_font and current_font not in allowed_fonts:
        fallback = get_fallback_font(current_font, spec)
        set_run_font(run, fallback)


# ──────────────────────────────────────────────────────────────
# 行距应用
# ──────────────────────────────────────────────────────────────
def apply_line_spacing(pf, line_spacing, style_line_spacing=None):
    """
    应用行距。
    style_line_spacing: 样式内嵌的行距（优先使用）
    line_spacing: 顶层行距配置（style_line_spacing 为空时使用）
    """
    ls = style_line_spacing if style_line_spacing else line_spacing
    if not ls:
        return

    ls_type = ls.get("type", "fixed")
    ls_value = ls.get("value", 28)  # 默认28pt固定行距

    pPr = pf._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(pf._element, qn("w:pPr"))
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn("w:spacing"))

    if ls_type == "fixed":
        # 固定行距（磅值 → twips：1pt = 20 twips）
        spacing.set(qn("w:line"), str(int(ls_value * 20)))
        spacing.set(qn("w:lineRule"), "exact")
    else:
        # 多倍行距（值是倍数，Word 用 240 表示 1.0 倍）
        spacing.set(qn("w:line"), str(int(ls_value * 240)))
        spacing.set(qn("w:lineRule"), "auto")


# ──────────────────────────────────────────────────────────────
# 格式应用（核心）
# ──────────────────────────────────────────────────────────────
def apply_style(para, style_name, spec, light_mode=False):
    """
    将指定样式应用到段落。
    light_mode=True：只修正字体/字号/行距，不重设对齐/缩进/间距
    """
    styles = spec.get("styles", {})
    if style_name not in styles:
        return

    s = styles[style_name]
    pf = para.paragraph_format
    allowed_fonts = spec.get("allowed_fonts", ["宋体", "黑体", "楷体_GB2312", "仿宋", "仿宋_GB2312"])

    # ── 空段落处理 ──
    # 对于 blank 类型，保持段落完全空白，不添加任何内容
    if style_name == "blank":
        # 清空所有 runs
        for run in para.runs:
            run.text = ""
        return

    # ── 分页标记处理 ──
    # 对于 page_break_marker 类型，在设置分页后清空内容
    if style_name == "page_break_marker":
        set_page_break_before(para)
        for run in para.runs:
            run.text = ""
        return

    # ── 获取/创建 run ──
    if para.runs:
        run = para.runs[0]
    else:
        # 不再使用零宽空格，改用空字符串
        run = para.add_run("")

    # ── 字体（全量模式 or light模式都处理）──
    if not light_mode:
        if "font" in s and s["font"]:
            set_run_font(run, s["font"])
            for r in para.runs[1:]:
                set_run_font(r, s["font"])
    else:
        check_and_fix_font(run, allowed_fonts, spec)
        for r in para.runs[1:]:
            check_and_fix_font(r, allowed_fonts, spec)

    # ── 字号 ──
    if not light_mode:
        if "font_size" in s and s["font_size"]:
            sz = Pt(s["font_size"])
            run.font.size = sz
            for r in para.runs[1:]:
                r.font.size = sz
    else:
        for r in para.runs:
            if r.font.size and (r.font.size < Pt(9) or r.font.size > Pt(22)):
                r.font.size = Pt(s.get("font_size", 14))

    # ── 加粗 ──
    if not light_mode:
        if "bold" in s and s["bold"] is not None:
            run.font.bold = s["bold"]
            for r in para.runs[1:]:
                r.font.bold = s["bold"]

    # ── 对齐（light模式跳过）──
    if not light_mode and "align" in s and s["align"]:
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        pf.alignment = align_map.get(s["align"], WD_ALIGN_PARAGRAPH.LEFT)

    # ── 间距（light模式跳过）──
    if not light_mode:
        if "space_before" in s:
            pf.space_before = Pt(s["space_before"])
        if "space_after" in s:
            pf.space_after = Pt(s["space_after"])

    # ── 行距 ──
    line_spacing = spec.get("line_spacing", {})
    style_ls = s.get("line_spacing", None)
    apply_line_spacing(pf, line_spacing, style_ls)

    # ── 首行缩进（light模式跳过）──
    if not light_mode:
        if "first_line_indent" in s:
            pf.first_line_indent = pt(s["first_line_indent"])

    # ── 左缩进（light模式跳过）──
    if not light_mode:
        if "left_indent" in s:
            pf.left_indent = pt(s["left_indent"])


def apply_page_setup_inner(section, spec):
    """设置单个 section 的页边距。"""
    ps = spec.get("page_setup", {})
    if not ps:
        return
    section.top_margin = cm(ps.get("top", 3.7))
    section.bottom_margin = cm(ps.get("bottom", 3.5))
    section.left_margin = cm(ps.get("left", 2.8))
    section.right_margin = cm(ps.get("right", 2.6))


# ──────────────────────────────────────────────────────────────
# 字体可用性检查
# ──────────────────────────────────────────────────────────────
FONT_CHECKED = False

def _check_font_availability(spec):
    """检查规范中指定的关键字体是否在系统中可用（仅 Linux fc-list 可用时）。"""
    global FONT_CHECKED
    if FONT_CHECKED:
        return
    FONT_CHECKED = True

    import subprocess
    import shutil
    styles = spec.get("styles", {})
    used_fonts = set()
    for sname, sspec in styles.items():
        fn = sspec.get("font", "")
        if fn:
            used_fonts.add(fn)

    # fc-list 仅 Linux 可用，macOS/Windows 跳过检测
    if not shutil.which("fc-list"):
        return

    try:
        result = subprocess.run(
            ["fc-list", ":lang=zh"],
            capture_output=True, text=True, timeout=5
        )
        installed = result.stdout.lower()
        for font in used_fonts:
            if font.lower() not in installed:
                print(f"⚠️ 字体未在系统中检测到: {font}")
                print(f"   文档会 fallback 到默认字体，建议安装")
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass


# ──────────────────────────────────────────────────────────────
# 主函数
# ──────────────────────────────────────────────────────────────
def format_docx(input_path, output_path=None, dry_run=False, light_mode=False):
    """
    主函数：对 DOCX 文件应用全局格式规范。
    """
    input_path = Path(input_path)
    if not input_path.exists():
        print(f"ERROR: 文件不存在: {input_path}")
        return False

    output_path = Path(output_path) if output_path else input_path

    spec = load_format_spec()

    mode_str = "light（轻量）" if light_mode else "full（全量）"
    if dry_run:
        print(f"[DRY RUN] 格式化（{mode_str}）: {input_path}")
    else:
        print(f"📋 使用格式规范 v{spec.get('version', '?')}（{mode_str}模式）")

    global PARA_IDX_COUNTER
    PARA_IDX_COUNTER = [0]

    doc = Document(str(input_path))
    total_paragraphs = len(doc.paragraphs)

    # 1. 页边距（覆盖所有 section）
    if not light_mode:
        for section in doc.sections:
            apply_page_setup_inner(section, spec)

    # 2. 字体可用性检查
    _check_font_availability(spec)

    # 3. 段落格式化
    stats = {}
    page_breaks_inserted = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        ptype = detect_type(para, spec)

        if not dry_run:
            apply_style(para, ptype, spec, light_mode=light_mode)

            # 3.5 分页控制 - 检查显式分页标记
            if should_insert_page_break(text, spec):
                set_page_break_before(para)
                page_breaks_inserted.append(idx)

        stats[ptype] = stats.get(ptype, 0) + 1

    # 4. 表格内单元格格式化
    PARA_IDX_COUNTER[0] = 100  # 跳过封面检测
    table_stats = {}
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    if row_idx == 0:
                        ptype = "table_header"
                    else:
                        ptype = "table_body"
                    if not dry_run:
                        apply_style(para, ptype, spec, light_mode=light_mode)
                    table_stats[ptype] = table_stats.get(ptype, 0) + 1

    if dry_run:
        total = sum(stats.values()) + sum(table_stats.values())
        print(f"[DRY RUN] 页边距(覆盖{len(doc.sections)}个section) + {total} 个段落已检测，格式未写入")
        print(f"[DRY RUN] 段落类型分布: {stats}")
        print(f"[DRY RUN] 表格段落类型分布: {table_stats}")
        return True

    # 保存
    tmp = output_path.with_suffix(".tmp.docx")
    doc.save(str(tmp))
    tmp.replace(output_path)

    print(f"✅ 格式化完成: {output_path}")
    print(f"   段落类型分布（按出现次数排序）: {dict(sorted(stats.items(), key=lambda x: -x[1]))}")
    print(f"   表格段落: {table_stats}")
    if page_breaks_inserted:
        print(f"   分页插入位置: {page_breaks_inserted}")
    return True


# ──────────────────────────────────────────────────────────────
# 入口
# ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_file = sys.argv[1]

    output_file = None
    dry_run = False
    light_mode = False

    for i, arg in enumerate(sys.argv[2:], start=2):
        if arg == "--dry-run":
            dry_run = True
        elif arg == "--light":
            light_mode = True
        elif output_file is None and not arg.startswith("--"):
            output_file = arg

    success = format_docx(input_file, output_file, dry_run=dry_run, light_mode=light_mode)
    sys.exit(0 if success else 1)